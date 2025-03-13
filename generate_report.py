import json
import re

import jinja2
import openai
from docx import Document


def read_docx(file_path: str):
    """
    读取docx 文档获取内容
    :param file_path: docx文件地址
    :return:
    """
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return ''.join(full_text)


def value_conversion(value: str, type: str):
    """
    结果转换器
    :param value:
    :param type:
    :return:
    """
    result = None
    if type == 'int':
        result = int(value)
    elif type == 'float':
        result = float(value)

    return result


# 提取所有标识符
data = read_docx('比亚迪汽车金融优先公司年度信息披露报告.docx')
targets = re.findall(pattern=r'(?<=\{{\s)(.+?)\s(?=\})', string=data)

with open('content_mapping.json', 'r', encoding='utf-8') as f:
    all_identifier = json.load(f)

# 读取文档信息
with open('byd_info.json', 'r', encoding='utf-8') as f:
    doc_contents = json.load(f)
doc_contents = doc_contents[0]

page_number = 1
for content in doc_contents['content']:
    del content['summary']
    del content['index']
    content['page_number'] = page_number
    page_number += 1
identifier_dict = {}

system_prompt = """
    # Role：数据解析师

    # Profile：
    - version：1.0
    - language：zh-cn
    - description：你将获得一个JSON数组作为现有数据，内容包含多个文件每页内容总结。你需要理解用户给出的操作提示，分析并通过对应操作如实准确无误的返回用户所需指定格式数据
    JSON结构：[
        {
            "filename": "文件名",
            "content": [{"description": "每一页文档的指标描述总结", "page_number": "第几页文档"}]
        }
    ]

    # Skills
    - 擅长分析理解用户提供的
    - 擅长

    # Rules

    # Workflow
"""
remote_llm = openai.OpenAI(
    api_key='sk-3fb76d31383b4552b9c3ebf82f44157d',
    base_url='https://dashscope.aliyuncs.com/compatible-mode/v1'
)
completion = remote_llm.chat.completions.create(
    temperature=0.6,
    model='qwen2.5-14b-instruct',
    messages=[
        {'role': 'system', 'content': system_prompt},
        {'role': 'user', 'content': '从第三页文档中提取`交易性金融資產`，只输出不带有逗号分隔符数值金额'},
        {'role': 'assistant', 'content': '25415146000'},
        {'role': 'user', 'content': user_question}
    ],
    max_tokens=8192,
    timeout=12000
)

remote_llm = openai.OpenAI(
    api_key='sk-3fb76d31383b4552b9c3ebf82f44157d',
    base_url='https://dashscope.aliyuncs.com/compatible-mode/v1'
)

# 调用远程llm 获取对应数据
for identifier, values in all_identifier.items():
    operate = values['operate']

    completion = remote_llm.chat.completions.create(
        temperature=0.6,
        model='qwen2.5-14b-instruct',
        messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': '从第三页文档中提取`交易性金融資產`，只输出不带有逗号分隔符数值金额'},
            {'role': 'assistant', 'content': '25415146000'},
            {'role': 'user', 'content': user_question}
        ],
        max_tokens=8192,
        timeout=12000
    )
    value = completion.choices[0].message.content
    value = value.strip()

    identifier_dict[identifier] = value

template = jinja2.Template(data)
data = template.render(**identifier_dict)
print(data)
